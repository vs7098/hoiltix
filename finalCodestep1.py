from PyPDF2 import PdfFileReader
from pdf2image import convert_from_path
from pyzbar.pyzbar import decode
from PIL import Image, UnidentifiedImageError
from docx import Document
from docx.shared import Inches, RGBColor, Pt  # Added Pt for font size
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_serial_background_color(run, color):
    """
    Apply background color to the serial number text.
    """
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)  # Use 'FF0000' for red background
    rPr.append(shd)

def set_serial_number_properties(run, color, font_size):
    """
    Set the serial number text properties (color, background, and font size).
    """
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    run.font.size = Pt(font_size)  # Set smaller font size
    set_serial_background_color(run, 'FF0000')  # Red background

def ensure_image_format(image_path):
    """
    Ensure the image is in a supported format for python-docx (PNG, JPEG, BMP, GIF).
    If not, convert it to PNG.
    """
    try:
        img = Image.open(image_path)
        if img.format not in ['JPEG', 'PNG', 'BMP', 'GIF']:
            # Convert the image to PNG if it's not in a supported format
            new_image_path = f"{os.path.splitext(image_path)[0]}.png"
            img.save(new_image_path, format='PNG')
            print(f"Image converted to {new_image_path}")
            return new_image_path
        return image_path
    except UnidentifiedImageError:
        print(f"Error: Unrecognized image format for {image_path}. Please check the file.")
        return None

def set_custom_page_size(doc, width, height, margin):
    """
    Set the custom page size and margins for the Word document.
    """
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(width)
        section.page_height = Inches(height)
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)

def extract_qr_images_to_word(pdf_path, word_path, input_image_path):
    # Ensure the image is in a supported format
    input_image_path = ensure_image_format(input_image_path)
    if input_image_path is None:
        return  # Exit if the image cannot be recognized or converted

    # Convert PDF pages to images
    images = convert_from_path(pdf_path, dpi=300)
    
    # Initialize Word document
    doc = Document()

    # Set custom page size and margins
    set_custom_page_size(doc, 6.50, 1.80, 0.1)

    # Define dimensions
    input_image_width = Inches(2.94)  # Image width for column 1
    qr_code_width = Inches(1.2)  # Smaller QR Code
    qr_code_height = Inches(1.2)  
    serial_number_width = Inches(1.2)
    row_height = Inches(1.72)  # Set a fixed height for the row
    serial_font_size = 8  # Set font size for Serial Number

    # Initialize serial number
    serial_prefix = "DESIMANDI"
    serial_number = 2001

    # Iterate through the images (one image per PDF page)
    for i, image in enumerate(images):
        # Extract QR code images
        decoded_objects = decode(image)
        
        # Create a table with 1 row and 2 columns
        table = doc.add_table(rows=1, cols=2)

        # Disable autofit to keep column widths fixed
        table.autofit = False

        # Set fixed column widths
        table.columns[0].width = Inches(2.94)  # First column for predefined image
        table.columns[1].width = Inches(2.5)  # Second column for QR code & serial number

        # Center the table
        table.alignment = 1  # Center alignment (1 = center)

        # Insert the predefined image into the first column
        cell_1 = table.cell(0, 0)
        paragraph_1 = cell_1.paragraphs[0]
        run_1 = paragraph_1.add_run()
        run_1.add_picture(input_image_path, width=input_image_width)

        # Insert QR code and Serial Number in the second column
        cell_2 = table.cell(0, 1)
        paragraph_2 = cell_2.paragraphs[0]

        # Handle QR code image
        qr_image_path = None
        if decoded_objects:
            for j, obj in enumerate(decoded_objects):
                qr_image = obj.rect
                qr_image_crop = image.crop((qr_image.left, qr_image.top, qr_image.left + qr_image.width, qr_image.top + qr_image.height))
                qr_image_path = f'qr_code_page_{i+1}_code_{j+1}.png'
                qr_image_crop.save(qr_image_path)
                break  # Take the first QR code found

        if qr_image_path and os.path.exists(qr_image_path):
            paragraph_2.add_run().add_picture(qr_image_path, width=qr_code_width, height=qr_code_height)
            os.remove(qr_image_path)  # Clean up temporary file
        else:
            paragraph_2.add_run("No QR code found")

        # Insert serial number directly below QR code
        paragraph_2.add_run("\n")  # Add newline for separation
        serial_run = paragraph_2.add_run(f"SR: {serial_prefix}{serial_number:05d}")
        set_serial_number_properties(serial_run, color='FFFFFF', font_size=serial_font_size)

        # Increment the serial number for the next page
        serial_number += 1001

        # Add a page break after each table to ensure each table is on a new page
        doc.add_page_break()

    # Save the Word document
    doc.save(word_path)
    print(f"QR codes and images extracted and saved to {word_path}")

# Example usage
pdf_file_path = 'dm2.pdf'
word_file_path = 'output_word_filedm2.docx'
input_image_path = '3.png'
extract_qr_images_to_word(pdf_file_path, word_file_path, input_image_path)

